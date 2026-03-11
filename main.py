import os
import json
import logging
from datetime import datetime, timedelta, date
from io import BytesIO
from contextlib import asynccontextmanager

from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, Depends, HTTPException, Query, Request
from fastapi.responses import HTMLResponse, StreamingResponse, FileResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from jose import JWTError, jwt
from sqlalchemy import text, func, select
from sqlalchemy.ext.asyncio import AsyncSession
import bcrypt
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from database import async_session, get_db, init_db, engine
from models import User, Document, CongVan
from search import (
    fulltext_search_documents, fulltext_search_cong_van,
    combined_search, semantic_search
)
from ai_extract import stream_chat_response
from crawler import run_crawl

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────
# Config
# ──────────────────────────────────────────────
JWT_SECRET = os.getenv("JWT_SECRET", "default-secret-change-me")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = int(os.getenv("JWT_EXPIRE_HOURS", "24"))
APP_VERSION = os.getenv("APP_VERSION", "1.0.0")


# ──────────────────────────────────────────────
# Lifespan
# ──────────────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    await seed_db()
    yield
    await engine.dispose()


app = FastAPI(title="VietTax Legal DB", version=APP_VERSION, lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ──────────────────────────────────────────────
# Auth helpers
# ──────────────────────────────────────────────
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()


def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())


def create_token(data: dict) -> str:
    expire = datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS)
    return jwt.encode({**data, "exp": expire}, JWT_SECRET, algorithm=JWT_ALGORITHM)


async def get_current_user(request: Request, db: AsyncSession = Depends(get_db)) -> User:
    token = request.headers.get("Authorization", "").replace("Bearer ", "")
    if not token:
        raise HTTPException(401, "Not authenticated")
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user_id = payload.get("sub")
        if not user_id:
            raise HTTPException(401, "Invalid token")
    except JWTError:
        raise HTTPException(401, "Invalid token")

    result = await db.execute(select(User).where(User.id == int(user_id)))
    user = result.scalar_one_or_none()
    if not user or not user.is_active:
        raise HTTPException(401, "User not found or inactive")
    return user


async def require_admin(user: User = Depends(get_current_user)) -> User:
    if user.role != "admin":
        raise HTTPException(403, "Admin access required")
    return user


# ──────────────────────────────────────────────
# Pydantic schemas
# ──────────────────────────────────────────────
class LoginRequest(BaseModel):
    username: str
    password: str

class RegisterRequest(BaseModel):
    username: str
    email: str
    password: str
    role: str = "user"

class ChatRequest(BaseModel):
    question: str

class ExportRequest(BaseModel):
    type: str  # search_results | document | chat_answer
    data: dict


# ──────────────────────────────────────────────
# Seed DB
# ──────────────────────────────────────────────
async def seed_db():
    """Import seed data if DB is empty. Also create default admin."""
    async with async_session() as db:
        # Create default admin
        result = await db.execute(select(User).where(User.username == "admin"))
        if not result.scalar_one_or_none():
            admin = User(
                username="admin",
                email="admin@vntaxdb.gpt4vn.com",
                hashed_password=hash_password("admin123"),
                role="admin"
            )
            db.add(admin)
            await db.commit()
            logger.info("Created default admin user")

        # Seed documents
        result = await db.execute(text("SELECT COUNT(*) FROM documents"))
        count = result.scalar()
        if count == 0:
            seed_file = os.path.join(os.path.dirname(__file__), "data", "seed_legal_db.json")
            if os.path.exists(seed_file):
                with open(seed_file, "r", encoding="utf-8") as f:
                    data = json.load(f)

                for item in data:
                    # Map _cat to sac_thue array
                    cat = item.pop("_cat", None)
                    sac_thue = item.pop("sac_thue", None)
                    if not sac_thue and cat:
                        sac_thue = [cat] if cat not in ("TaxAdmin", "DatDai", "HoKinhDoanh",
                                                         "TransferPricing", "InternationalTax") else []

                    # Map loai names to consistent format
                    loai_raw = item.get("loai", "")
                    loai = loai_raw
                    if "luật" in loai_raw.lower() and "sửa" not in loai_raw.lower():
                        loai = "Luật"
                    elif "nghị định" in loai_raw.lower():
                        loai = "NĐ"
                    elif "thông tư" in loai_raw.lower() and "thông tin" not in loai_raw.lower():
                        loai = "TT"
                    elif "quyết định" in loai_raw.lower():
                        loai = "QĐ"
                    elif "nghị quyết" in loai_raw.lower():
                        loai = "NQ"

                    # Handle sua_doi_boi as JSON string
                    sua_doi_boi = item.pop("sua_doi_boi", None)
                    if isinstance(sua_doi_boi, list):
                        sua_doi_boi = json.dumps(sua_doi_boi)

                    # Parse dates
                    ngay_ban_hanh = item.pop("ngay_ban_hanh", None)
                    hieu_luc_tu = item.pop("hieu_luc_tu", None)

                    def parse_date(d):
                        if not d:
                            return None
                        try:
                            return datetime.strptime(d, "%Y-%m-%d").date()
                        except (ValueError, TypeError):
                            return None

                    # Remove fields not in model
                    item.pop("luu_y_them", None)

                    doc = Document(
                        so_hieu=item.get("so_hieu", ""),
                        ten=item.get("ten", ""),
                        loai=loai,
                        co_quan=item.get("co_quan"),
                        ngay_ban_hanh=parse_date(ngay_ban_hanh),
                        hieu_luc_tu=parse_date(hieu_luc_tu),
                        tinh_trang=item.get("tinh_trang"),
                        sua_doi_boi=sua_doi_boi,
                        sac_thue=sac_thue if sac_thue else None,
                        tom_tat=item.get("tom_tat"),
                        link_tvpl=item.get("link_tvpl"),
                        link_vbpl=item.get("link_vbpl"),
                        tu_khoa=item.get("tu_khoa"),
                        luu_y=item.get("luu_y"),
                    )
                    db.add(doc)

                await db.commit()
                logger.info(f"Seeded {len(data)} documents from seed file")


# ──────────────────────────────────────────────
# Auth routes
# ──────────────────────────────────────────────
@app.post("/auth/login")
async def login(req: LoginRequest, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(User).where(User.username == req.username))
    user = result.scalar_one_or_none()
    if not user or not verify_password(req.password, user.hashed_password):
        raise HTTPException(401, "Invalid credentials")
    token = create_token({"sub": str(user.id), "role": user.role})
    return {"access_token": token, "token_type": "bearer", "role": user.role, "username": user.username}


@app.post("/auth/register")
async def register(req: RegisterRequest, db: AsyncSession = Depends(get_db),
                   admin: User = Depends(require_admin)):
    # Check existing
    result = await db.execute(select(User).where(
        (User.username == req.username) | (User.email == req.email)
    ))
    if result.scalar_one_or_none():
        raise HTTPException(400, "Username or email already exists")

    user = User(
        username=req.username,
        email=req.email,
        hashed_password=hash_password(req.password),
        role=req.role
    )
    db.add(user)
    await db.commit()
    return {"message": "User created", "id": user.id}


@app.get("/auth/me")
async def get_me(user: User = Depends(get_current_user)):
    return {"id": user.id, "username": user.username, "email": user.email, "role": user.role}


# ──────────────────────────────────────────────
# Documents API
# ──────────────────────────────────────────────
@app.get("/api/documents")
async def list_documents(
    q: str = "",
    loai: str = "",
    sac_thue: str = "",
    tinh_trang: str = "",
    year: str = "",
    limit: int = Query(50, le=200),
    offset: int = 0,
    db: AsyncSession = Depends(get_db)
):
    filters = {}
    if loai:
        filters["loai"] = loai
    if sac_thue:
        filters["sac_thue"] = sac_thue
    if tinh_trang:
        filters["tinh_trang"] = tinh_trang
    if year:
        filters["year"] = year

    result = await fulltext_search_documents(db, q, filters, limit, offset)

    # Serialize dates
    for item in result["items"]:
        for key in ("ngay_ban_hanh", "hieu_luc_tu"):
            if key in item and isinstance(item[key], date):
                item[key] = item[key].isoformat()

    return result


@app.get("/api/documents/{so_hieu:path}")
async def get_document(so_hieu: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.so_hieu == so_hieu))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    return {
        "id": doc.id,
        "so_hieu": doc.so_hieu,
        "ten": doc.ten,
        "loai": doc.loai,
        "co_quan": doc.co_quan,
        "ngay_ban_hanh": doc.ngay_ban_hanh.isoformat() if doc.ngay_ban_hanh else None,
        "hieu_luc_tu": doc.hieu_luc_tu.isoformat() if doc.hieu_luc_tu else None,
        "het_hieu_luc_tu": doc.het_hieu_luc_tu.isoformat() if doc.het_hieu_luc_tu else None,
        "tinh_trang": doc.tinh_trang,
        "thay_the_boi": doc.thay_the_boi,
        "sua_doi_boi": json.loads(doc.sua_doi_boi) if doc.sua_doi_boi else [],
        "sac_thue": doc.sac_thue or [],
        "category": doc.category or [],
        "tom_tat": doc.tom_tat,
        "noi_dung": doc.noi_dung,
        "link_tvpl": doc.link_tvpl,
        "link_vbpl": doc.link_vbpl,
        "tu_khoa": doc.tu_khoa or [],
        "luu_y": doc.luu_y,
    }


@app.get("/api/documents/{so_hieu:path}/relations")
async def get_document_relations(so_hieu: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(Document).where(Document.so_hieu == so_hieu))
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    relations = {"thay_the_boi": None, "sua_doi_boi": [], "duoc_tham_chieu_boi": []}

    if doc.thay_the_boi:
        r = await db.execute(select(Document).where(Document.so_hieu == doc.thay_the_boi))
        rep = r.scalar_one_or_none()
        if rep:
            relations["thay_the_boi"] = {"so_hieu": rep.so_hieu, "ten": rep.ten}

    if doc.sua_doi_boi:
        try:
            so_hieus = json.loads(doc.sua_doi_boi)
            for sh in so_hieus:
                r = await db.execute(select(Document).where(Document.so_hieu == sh))
                d = r.scalar_one_or_none()
                if d:
                    relations["sua_doi_boi"].append({"so_hieu": d.so_hieu, "ten": d.ten})
        except json.JSONDecodeError:
            pass

    # Find docs that reference this one
    r = await db.execute(
        select(Document).where(Document.sua_doi_boi.contains(so_hieu))
    )
    for d in r.scalars().all():
        relations["duoc_tham_chieu_boi"].append({"so_hieu": d.so_hieu, "ten": d.ten})

    return relations


# ──────────────────────────────────────────────
# Công văn API
# ──────────────────────────────────────────────
@app.get("/api/cong-van")
async def list_cong_van(
    q: str = "",
    co_quan: str = "",
    sac_thue: str = "",
    year: str = "",
    limit: int = Query(50, le=200),
    offset: int = 0,
    db: AsyncSession = Depends(get_db)
):
    filters = {}
    if co_quan:
        filters["co_quan"] = co_quan
    if sac_thue:
        filters["sac_thue"] = sac_thue
    if year:
        filters["year"] = year

    result = await fulltext_search_cong_van(db, q, filters, limit, offset)

    for item in result["items"]:
        for key in ("ngay_ban_hanh",):
            if key in item and isinstance(item[key], date):
                item[key] = item[key].isoformat()

    return result


@app.get("/api/cong-van/{so_hieu:path}")
async def get_cong_van(so_hieu: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(CongVan).where(CongVan.so_hieu == so_hieu))
    cv = result.scalar_one_or_none()
    if not cv:
        raise HTTPException(404, "Công văn not found")

    return {
        "id": cv.id,
        "so_hieu": cv.so_hieu,
        "ten": cv.ten,
        "co_quan": cv.co_quan,
        "nguoi_nhan": cv.nguoi_nhan,
        "ngay_ban_hanh": cv.ngay_ban_hanh.isoformat() if cv.ngay_ban_hanh else None,
        "sac_thue": cv.sac_thue or [],
        "van_ban_trich_dan": cv.van_ban_trich_dan,
        "ket_luan": cv.ket_luan,
        "noi_dung_day_du": cv.noi_dung_day_du,
        "tags": cv.tags or [],
        "link_tvpl": cv.link_tvpl,
    }


# ──────────────────────────────────────────────
# Search API
# ──────────────────────────────────────────────
@app.get("/api/search")
async def search(
    q: str = "",
    type: str = "all",
    sac_thue: str = "",
    loai: str = "",
    tinh_trang: str = "",
    year: str = "",
    limit: int = Query(50, le=200),
    offset: int = 0,
    db: AsyncSession = Depends(get_db)
):
    filters = {}
    if sac_thue:
        filters["sac_thue"] = sac_thue
    if loai:
        filters["loai"] = loai
    if tinh_trang:
        filters["tinh_trang"] = tinh_trang
    if year:
        filters["year"] = year

    result = await combined_search(db, q, type, filters, limit, offset)

    # Serialize dates
    for key in ("documents", "cong_van"):
        if key in result:
            for item in result[key].get("items", []):
                for dkey in ("ngay_ban_hanh", "hieu_luc_tu"):
                    if dkey in item and isinstance(item[dkey], date):
                        item[dkey] = item[dkey].isoformat()

    return result


# ──────────────────────────────────────────────
# AI Chatbox
# ──────────────────────────────────────────────
CHAT_SYSTEM_PROMPT = """Bạn là chuyên gia tư vấn thuế Việt Nam.
Dựa trên các văn bản pháp luật và công văn được cung cấp, hãy trả lời câu hỏi.

Format trả lời BẮT BUỘC:
## 1. Quy định áp dụng
[Trích dẫn cụ thể: điều/khoản/mục của Luật/NĐ/TT, ghi rõ còn hiệu lực hay không]

## 2. Công văn hướng dẫn liên quan
[Số CV, cơ quan ban hành, ngày, kết luận chính]

## 3. Rủi ro & lưu ý
[Các trường hợp không được áp dụng, điều kiện cần đáp ứng]

## 4. Đề xuất
[Hướng xử lý thực tế]

QUAN TRỌNG: Chỉ trích dẫn văn bản có trong context. Ghi rõ nếu không tìm thấy quy định cụ thể."""


def build_legal_context(docs: list[dict], cvs: list[dict]) -> str:
    """Build context string from search results."""
    parts = []

    if docs:
        parts.append("=== VĂN BẢN PHÁP LUẬT ===")
        for d in docs:
            parts.append(f"\n--- {d.get('so_hieu', '')} ---")
            parts.append(f"Tên: {d.get('ten', '')}")
            parts.append(f"Loại: {d.get('loai', '')} | Tình trạng: {d.get('tinh_trang', '')}")
            if d.get("tom_tat"):
                parts.append(f"Tóm tắt: {d['tom_tat']}")
            if d.get("luu_y"):
                parts.append(f"Lưu ý: {d['luu_y']}")
            if d.get("noi_dung"):
                parts.append(f"Nội dung: {d['noi_dung'][:3000]}")

    if cvs:
        parts.append("\n=== CÔNG VĂN ===")
        for c in cvs:
            parts.append(f"\n--- {c.get('so_hieu', '')} ---")
            if c.get("ten"):
                parts.append(f"Tên: {c['ten']}")
            parts.append(f"Cơ quan: {c.get('co_quan', '')}")
            if c.get("ket_luan"):
                parts.append(f"Kết luận: {c['ket_luan']}")
            if c.get("noi_dung_day_du"):
                parts.append(f"Nội dung: {c['noi_dung_day_du'][:3000]}")

    return "\n".join(parts)


@app.post("/api/chat")
async def chat(req: ChatRequest, db: AsyncSession = Depends(get_db)):
    question = req.question.strip()
    if not question:
        raise HTTPException(400, "Question is required")

    # Search for relevant context
    relevant_docs = await semantic_search(db, question, "documents", limit=5)
    relevant_cvs = await semantic_search(db, question, "cong_van", limit=5)

    # Fallback to fulltext if no semantic results
    if not relevant_docs:
        ft_result = await fulltext_search_documents(db, question, {}, limit=5)
        relevant_docs = ft_result["items"]

    context = build_legal_context(relevant_docs, relevant_cvs)

    def generate():
        for chunk in stream_chat_response(CHAT_SYSTEM_PROMPT, context, question):
            yield f"data: {json.dumps({'text': chunk}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


# ──────────────────────────────────────────────
# Export Word
# ──────────────────────────────────────────────
def create_word_doc(title: str, content: str) -> BytesIO:
    """Create a Word document with given content."""
    doc = DocxDocument()

    # Title
    heading = doc.add_heading(title, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(2, 138, 57)  # #028a39

    # Metadata
    doc.add_paragraph(f"Xuất từ VietTax Legal DB — {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")

    # Content — split by lines and handle markdown headings
    for line in content.split("\n"):
        line = line.strip()
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.post("/api/export/word")
async def export_word(req: ExportRequest):
    data = req.data

    if req.type == "document":
        title = f"{data.get('so_hieu', '')} — {data.get('ten', '')}"
        lines = []
        if data.get("tom_tat"):
            lines.append(f"Tóm tắt: {data['tom_tat']}")
        if data.get("tinh_trang"):
            lines.append(f"Tình trạng: {data['tinh_trang']}")
        if data.get("noi_dung"):
            lines.append(f"\n{data['noi_dung']}")
        content = "\n".join(lines)

    elif req.type == "chat_answer":
        title = f"Tư vấn thuế: {data.get('question', '')[:100]}"
        content = data.get("answer", "")

    elif req.type == "search_results":
        title = f"Kết quả tìm kiếm: {data.get('query', '')}"
        lines = []
        for i, item in enumerate(data.get("items", [])[:50], 1):
            lines.append(f"{i}. {item.get('so_hieu', '')} — {item.get('ten', '')}")
            if item.get("tom_tat"):
                lines.append(f"   {item['tom_tat']}")
            lines.append("")
        content = "\n".join(lines)
    else:
        raise HTTPException(400, "Invalid export type")

    buffer = create_word_doc(title, content)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=viettax-export-{datetime.now().strftime('%Y%m%d-%H%M%S')}.docx"}
    )


# ──────────────────────────────────────────────
# Admin API
# ──────────────────────────────────────────────
@app.get("/api/admin/stats")
async def admin_stats(db: AsyncSession = Depends(get_db), admin: User = Depends(require_admin)):
    docs = await db.execute(text("SELECT COUNT(*) FROM documents"))
    cvs = await db.execute(text("SELECT COUNT(*) FROM cong_van"))
    users = await db.execute(text("SELECT COUNT(*) FROM users"))

    return {
        "documents": docs.scalar(),
        "cong_van": cvs.scalar(),
        "users": users.scalar(),
    }


@app.post("/api/admin/crawl")
async def admin_crawl(
    loai: str = Query("Luật"),
    max_pages: int = Query(3, le=10),
    db: AsyncSession = Depends(get_db),
    admin: User = Depends(require_admin)
):
    results = await run_crawl(loai=loai, max_pages=max_pages)
    imported = 0

    for item in results:
        existing = await db.execute(select(Document).where(Document.so_hieu == item["so_hieu"]))
        if existing.scalar_one_or_none():
            continue

        doc = Document(
            so_hieu=item["so_hieu"],
            ten=item.get("ten", ""),
            loai=loai,
            co_quan=item.get("co_quan"),
            ngay_ban_hanh=datetime.strptime(item["ngay_ban_hanh"], "%Y-%m-%d").date() if item.get("ngay_ban_hanh") else None,
            hieu_luc_tu=datetime.strptime(item["hieu_luc_tu"], "%Y-%m-%d").date() if item.get("hieu_luc_tu") else None,
            tinh_trang=item.get("tinh_trang"),
            noi_dung=item.get("noi_dung"),
            link_vbpl=item.get("link_vbpl"),
        )
        db.add(doc)
        imported += 1

    await db.commit()
    return {"crawled": len(results), "imported": imported}


@app.get("/api/admin/users")
async def admin_list_users(db: AsyncSession = Depends(get_db), admin: User = Depends(require_admin)):
    result = await db.execute(select(User).order_by(User.created_at.desc()))
    users = result.scalars().all()
    return [
        {"id": u.id, "username": u.username, "email": u.email, "role": u.role,
         "is_active": u.is_active, "created_at": u.created_at.isoformat() if u.created_at else None}
        for u in users
    ]


@app.post("/api/admin/users")
async def admin_create_user(req: RegisterRequest, db: AsyncSession = Depends(get_db),
                             admin: User = Depends(require_admin)):
    result = await db.execute(select(User).where(
        (User.username == req.username) | (User.email == req.email)
    ))
    if result.scalar_one_or_none():
        raise HTTPException(400, "Username or email already exists")

    user = User(
        username=req.username,
        email=req.email,
        hashed_password=hash_password(req.password),
        role=req.role
    )
    db.add(user)
    await db.commit()
    return {"message": "User created", "id": user.id}


# ──────────────────────────────────────────────
# Serve frontend
# ──────────────────────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    html_path = os.path.join(os.path.dirname(__file__), "static", "index.html")
    if os.path.exists(html_path):
        with open(html_path, "r", encoding="utf-8") as f:
            return HTMLResponse(f.read())
    return HTMLResponse("<h1>VietTax Legal DB</h1><p>Frontend not found.</p>")


# Health check
@app.get("/health")
async def health():
    return {"status": "ok", "version": APP_VERSION}
